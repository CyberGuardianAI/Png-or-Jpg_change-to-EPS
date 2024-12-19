import tkinter as tk
from tkinter import filedialog, messagebox
from googletrans import Translator
from docx import Document
from pdfminer.high_level import extract_text
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# 创建翻译函数
def translate_text(text, translation_direction):
    translator = Translator()

    try:
        if translation_direction == "en2zh":
            return translator.translate(text, src='en', dest='zh-cn').text
        elif translation_direction == "zh2en":
            return translator.translate(text, src='zh-cn', dest='en').text
    except Exception as e:
        return f"翻译失败: {e}"


# 解析PDF文件，提取文本
def extract_text_from_pdf(pdf_path):
    try:
        text = extract_text(pdf_path)
        return text
    except Exception as e:
        messagebox.showerror("错误", f"PDF文件提取失败: {e}")
        return None


# 解析Word文件，提取文本
def extract_text_from_word(word_path):
    try:
        doc = Document(word_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)
    except Exception as e:
        messagebox.showerror("错误", f"Word文件提取失败: {e}")
        return None


# 处理并翻译PDF文件
def process_pdf(pdf_path, output_path, translation_direction):
    text = extract_text_from_pdf(pdf_path)
    if text:
        translated_text = translate_text(text, translation_direction)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(translated_text)
        messagebox.showinfo("成功", f"翻译完成！文件已保存为 {output_path}")


# 处理并翻译Word文件
def process_word(word_path, output_path, translation_direction):
    text = extract_text_from_word(word_path)
    if text:
        translated_text = translate_text(text, translation_direction)

        # 生成翻译后的Word文档
        doc = Document()
        translated_paragraphs = translated_text.split("\n")

        for para in translated_paragraphs:
            if para.strip():  # 如果段落有内容
                doc.add_paragraph(para)

        doc.save(output_path)
        messagebox.showinfo("成功", f"翻译完成！文件已保存为 {output_path}")


# 选择文件函数
def select_input_file():
    file_selected = filedialog.askopenfilename(filetypes=[("PDF文件", "*.pdf"), ("Word文件", "*.docx")])
    if file_selected:
        input_file_var.set(file_selected)


# 选择输出文件函数
def select_output_file():
    file_selected = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word文件", "*.docx")])
    if file_selected:
        output_file_var.set(file_selected)


# 翻译文档的主函数
def translate_document():
    input_file = input_file_var.get()
    output_file = output_file_var.get()
    translation_direction = lang_var.get()

    if not input_file or not output_file:
        messagebox.showerror("错误", "请输入有效的文件路径")
        return

    try:
        # 根据文件类型进行不同的处理
        if input_file.endswith(".pdf"):
            process_pdf(input_file, output_file, translation_direction)
        elif input_file.endswith(".docx"):
            process_word(input_file, output_file, translation_direction)
        else:
            messagebox.showerror("错误", "不支持的文件格式")
    except Exception as e:
        messagebox.showerror("错误", f"翻译过程出错: {e}")


# 创建主界面
root = tk.Tk()
root.title("自动翻译文档")

# 设置界面尺寸和背景颜色
root.geometry("600x400")
root.config(bg="#f0f0f0")

# 语言选择框
lang_var = tk.StringVar()
lang_var.set("en2zh")  # 默认翻译方向为英文到中文

lang_label = tk.Label(root, text="选择翻译方向：", font=("Arial", 12), bg="#f0f0f0")
lang_label.pack(pady=10)

lang_en2zh_button = tk.Radiobutton(root, text="英文 -> 中文", variable=lang_var, value="en2zh", font=("Arial", 12),
                                   bg="#f0f0f0")
lang_en2zh_button.pack()

lang_zh2en_button = tk.Radiobutton(root, text="中文 -> 英文", variable=lang_var, value="zh2en", font=("Arial", 12),
                                   bg="#f0f0f0")
lang_zh2en_button.pack()

# 文件选择按钮
select_file_button = tk.Button(root, text="选择PDF或Word文件", command=select_input_file, width=20, height=2,
                               font=("Arial", 12), bg="#4CAF50", fg="white", relief="raised")
select_file_button.pack(pady=10)

# 输入文件路径
input_file_var = tk.StringVar()
input_file_entry = tk.Entry(root, textvariable=input_file_var, width=50, font=("Arial", 12), state="readonly")
input_file_entry.pack(pady=5)

# 输出文件路径
output_file_var = tk.StringVar()
output_file_entry = tk.Entry(root, textvariable=output_file_var, width=50, font=("Arial", 12), state="readonly")
output_file_entry.pack(pady=5)

# 选择输出文件按钮
select_output_button = tk.Button(root, text="选择输出文件", command=select_output_file, width=20, height=2,
                                 font=("Arial", 12), bg="#4CAF50", fg="white", relief="raised")
select_output_button.pack(pady=10)

# 翻译按钮
translate_button = tk.Button(root, text="开始翻译", command=translate_document, width=20, height=2, font=("Arial", 12),
                             bg="#007bff", fg="white", relief="raised")
translate_button.pack(pady=20)

# 运行界面
root.mainloop()
